from fastapi import FastAPI, APIRouter, HTTPException, UploadFile, File, Form, Depends, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from dotenv import load_dotenv
from starlette.middleware.cors import CORSMiddleware
from motor.motor_asyncio import AsyncIOMotorClient
import os
import logging
from pathlib import Path
from pydantic import BaseModel, Field, EmailStr
from typing import List, Optional, Dict, Any
import uuid
from datetime import datetime, timedelta
import asyncio
import aiohttp
import json
import jwt
import bcrypt
from fastapi.responses import JSONResponse
import PyPDF2
import pdfplumber
from docx import Document
import io

# Load environment variables
ROOT_DIR = Path(__file__).parent
load_dotenv(ROOT_DIR / '.env')

# MongoDB connection (disabled for now)
# mongo_url = os.environ['MONGO_URL']
# client = AsyncIOMotorClient(mongo_url)
# db = client[os.environ['DB_NAME']]

# JWT Configuration
JWT_SECRET = os.environ.get('JWT_SECRET', 'your-secret-key-here')
JWT_ALGORITHM = os.environ.get('JWT_ALGORITHM', 'HS256')
JWT_EXPIRY_MINUTES = int(os.environ.get('JWT_EXPIRY_MINUTES', '1440'))

# Create the main app without a prefix
app = FastAPI()

# Create a router with the /api prefix
api_router = APIRouter(prefix="/api")

# Security
security = HTTPBearer(auto_error=False)

# Gemini API configuration
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY', 'AIzaSyDwh7j95PFsTaxM_UbU_adrGbiDwp3UdK8')
GEMINI_API_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"

# Define Models
class StatusCheck(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    client_name: str
    timestamp: datetime = Field(default_factory=datetime.utcnow)

class StatusCheckCreate(BaseModel):
    client_name: str

class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str

class UserLogin(BaseModel):
    email: EmailStr
    password: str

class User(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    email: EmailStr
    name: str
    hashed_password: str
    created_at: datetime = Field(default_factory=datetime.utcnow)
    is_active: bool = True

class UserResponse(BaseModel):
    id: str
    email: str
    name: str
    created_at: datetime
    is_active: bool

class TokenResponse(BaseModel):
    access_token: str
    token_type: str
    user: UserResponse

class DocumentAnalysisRequest(BaseModel):
    document_content: str
    analysis_type: str = "gap_analysis"
    
class DocumentAnalysisResponse(BaseModel):
    id: str
    analysis_result: Dict[str, Any]
    document_length: int
    analysis_type: str
    timestamp: datetime
    user_id: Optional[str] = None
    
class DocumentAnalysis(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_content: str
    analysis_result: Dict[str, Any]
    document_length: int
    analysis_type: str
    timestamp: datetime = Field(default_factory=datetime.utcnow)
    user_id: Optional[str] = None

class WorkflowRequest(BaseModel):
    document_content: str
    workflow_type: str = "user_journey"  # user_journey, service_blueprint, feature_flow

class WorkflowNode(BaseModel):
    id: str
    type: str  # start, process, decision, end
    label: str
    x: int
    y: int
    connections: List[str] = []

class WorkflowResponse(BaseModel):
    id: str
    workflow_nodes: List[WorkflowNode]
    workflow_type: str
    document_length: int
    timestamp: datetime
    user_id: Optional[str] = None

class Workflow(BaseModel):
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_content: str
    workflow_nodes: List[WorkflowNode]
    workflow_type: str
    document_length: int
    timestamp: datetime = Field(default_factory=datetime.utcnow)
    user_id: Optional[str] = None

# Document Parser Service
class DocumentParser:
    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """Parse PDF content using multiple methods"""
        try:
            # Method 1: Try pdfplumber first (better text extraction)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                if text.strip():
                    return text.strip()
        except Exception as e:
            logging.warning(f"pdfplumber failed: {e}")
        
        try:
            # Method 2: Fallback to PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            if text.strip():
                return text.strip()
        except Exception as e:
            logging.warning(f"PyPDF2 failed: {e}")
        
        raise ValueError("Could not extract text from PDF")
    
    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """Parse DOCX content"""
        try:
            doc = Document(io.BytesIO(file_content))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise ValueError(f"Could not extract text from DOCX: {str(e)}")
    
    @staticmethod
    def parse_text(file_content: bytes) -> str:
        """Parse plain text content"""
        try:
            return file_content.decode('utf-8').strip()
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return file_content.decode(encoding).strip()
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file")

# Authentication Service
class AuthService:
    @staticmethod
    def hash_password(password: str) -> str:
        """Hash password using bcrypt"""
        return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')
    
    @staticmethod
    def verify_password(plain_password: str, hashed_password: str) -> bool:
        """Verify password against hash"""
        return bcrypt.checkpw(plain_password.encode('utf-8'), hashed_password.encode('utf-8'))
    
    @staticmethod
    def create_access_token(data: dict) -> str:
        """Create JWT access token"""
        to_encode = data.copy()
        expire = datetime.utcnow() + timedelta(minutes=JWT_EXPIRY_MINUTES)
        to_encode.update({"exp": expire})
        return jwt.encode(to_encode, JWT_SECRET, algorithm=JWT_ALGORITHM)
    
    @staticmethod
    def verify_token(token: str) -> dict:
        """Verify JWT token"""
        try:
            payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
            return payload
        except jwt.ExpiredSignatureError:
            raise HTTPException(status_code=401, detail="Token has expired")
        except jwt.JWTError:
            raise HTTPException(status_code=401, detail="Invalid token")

# Dependency to get current user
async def get_current_user(credentials: HTTPAuthorizationCredentials = Depends(security)) -> Optional[User]:
    """Get current user from JWT token"""
    try:
        payload = AuthService.verify_token(credentials.credentials)
        user_id = payload.get("user_id")
        if not user_id:
            raise HTTPException(status_code=401, detail="Invalid token")
        
        raise HTTPException(status_code=503, detail="Database is disabled. User lookup unavailable.")
    except Exception as e:
        raise HTTPException(status_code=401, detail=str(e))

# Optional dependency for endpoints that work with or without auth
async def get_current_user_optional(credentials: HTTPAuthorizationCredentials = Depends(security)) -> Optional[User]:
    """Get current user from JWT token, but don't require it"""
    try:
        if not credentials:
            return None
        return await get_current_user(credentials)
    except:
        return None

# Gemini AI Service
class GeminiAIService:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = GEMINI_API_URL
        self.max_retries = 3
        self.base_delay = 2  # seconds
        
    async def analyze_document(self, document_content: str, analysis_type: str) -> Dict[str, Any]:
        """Analyze document using Gemini AI with retry logic"""
        
        # Create system prompts for different analysis types
        system_prompts = {
            "gap_analysis": """You are an expert business analyst and product manager. Analyze the provided PRD/BRD document and identify:

1. BUSINESS GAPS: Areas where business goals are not clearly addressed by product requirements
2. DESIGN AMBIGUITIES: Vague or unclear design specifications that need clarification
3. MISSING REQUIREMENTS: Important functional or non-functional requirements that are missing
4. EDGE CASES: Potential scenarios, boundary conditions, or exceptional cases not covered
5. RECOMMENDATIONS: Specific actionable suggestions to improve the document

Please provide a structured analysis in JSON format with the following structure:
{
  "business_gaps": ["List of specific business gaps"],
  "design_ambiguities": ["List of design ambiguities"],
  "missing_requirements": ["List of missing requirements"],
  "edge_cases": ["List of edge cases"],
  "recommendations": ["List of specific recommendations"],
  "overall_assessment": "Overall assessment of the document quality and completeness"
}

Make sure your response is valid JSON and provides specific, actionable insights.""",
            
            "requirements_extraction": """You are an expert business analyst. Extract and categorize all requirements from the provided PRD/BRD document.

Please provide a structured extraction in JSON format:
{
  "functional_requirements": ["List of functional requirements"],
  "non_functional_requirements": ["List of non-functional requirements"],
  "business_requirements": ["List of business requirements"],
  "user_stories": ["List of user stories"],
  "acceptance_criteria": ["List of acceptance criteria"],
  "constraints": ["List of constraints"],
  "assumptions": ["List of assumptions"]
}""",
            
            "summary": """You are an expert document summarizer. Provide a comprehensive summary of the PRD/BRD document.

Please provide a structured summary in JSON format:
{
  "executive_summary": "Brief executive summary",
  "key_features": ["List of key features"],
  "target_audience": "Target audience description",
  "business_goals": ["List of business goals"],
  "success_metrics": ["List of success metrics"],
  "timeline": "Project timeline information",
  "stakeholders": ["List of stakeholders"],
  "risks": ["List of identified risks"]
}"""
        }
        
        prompt = system_prompts.get(analysis_type, system_prompts["gap_analysis"])
        
        payload = {
            "contents": [
                {
                    "parts": [
                        {
                            "text": f"{prompt}\n\nDocument to analyze:\n\n{document_content}"
                        }
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.3,
                "topK": 40,
                "topP": 0.95,
                "maxOutputTokens": 2048,
            }
        }
        
        headers = {
            "Content-Type": "application/json",
            "X-goog-api-key": self.api_key
        }
        
        # Retry logic with exponential backoff
        for attempt in range(self.max_retries):
            try:
                timeout_duration = 30 + (attempt * 15)  # Increase timeout with retries
                
                async with aiohttp.ClientSession() as session:
                    async with session.post(
                        self.base_url, 
                        json=payload, 
                        headers=headers,
                        timeout=aiohttp.ClientTimeout(total=timeout_duration)
                    ) as response:
                        response_text = await response.text()
                        
                        if response.status == 200:
                            result = json.loads(response_text)
                            
                            # Extract the generated content
                            if 'candidates' in result and len(result['candidates']) > 0:
                                content = result['candidates'][0]['content']['parts'][0]['text']
                                
                                # Try to parse as JSON
                                try:
                                    # Clean up the content first - remove markdown code blocks
                                    cleaned_content = content.strip()
                                    if cleaned_content.startswith('```json'):
                                        cleaned_content = cleaned_content.replace('```json', '').replace('```', '').strip()
                                    elif cleaned_content.startswith('```'):
                                        cleaned_content = cleaned_content.replace('```', '').strip()
                                    
                                    parsed_content = json.loads(cleaned_content)
                                    return parsed_content
                                except json.JSONDecodeError:
                                    # If not valid JSON, return as raw text but try to extract any JSON within
                                    import re
                                    json_match = re.search(r'\{.*\}', content, re.DOTALL)
                                    if json_match:
                                        try:
                                            parsed_content = json.loads(json_match.group())
                                            return parsed_content
                                        except json.JSONDecodeError:
                                            pass
                                    return {"raw_analysis": content}
                            else:
                                raise HTTPException(status_code=500, detail="No content generated by Gemini")
                        
                        elif response.status == 503:
                            # Service overloaded - retry with backoff
                            if attempt < self.max_retries - 1:
                                delay = self.base_delay * (2 ** attempt)  # Exponential backoff
                                logging.warning(f"Gemini API overloaded (503), retrying in {delay}s (attempt {attempt + 1}/{self.max_retries})")
                                await asyncio.sleep(delay)
                                continue
                            else:
                                # Final attempt failed
                                raise HTTPException(
                                    status_code=503, 
                                    detail="Gemini AI service is currently overloaded. This is a temporary issue with high demand. Please try again in a few minutes. Our system will automatically retry failed requests."
                                )
                        
                        elif response.status == 429:
                            # Rate limited - retry with longer delay
                            if attempt < self.max_retries - 1:
                                delay = self.base_delay * (3 ** attempt)  # Longer delay for rate limits
                                logging.warning(f"Gemini API rate limited (429), retrying in {delay}s (attempt {attempt + 1}/{self.max_retries})")
                                await asyncio.sleep(delay)
                                continue
                            else:
                                raise HTTPException(
                                    status_code=429,
                                    detail="Rate limit exceeded. Please wait a moment before trying again."
                                )
                        
                        else:
                            # Other errors - don't retry
                            try:
                                error_data = json.loads(response_text)
                                error_message = error_data.get('error', {}).get('message', 'Unknown error')
                                raise HTTPException(
                                    status_code=response.status, 
                                    detail=f"Gemini API error: {error_message}"
                                )
                            except json.JSONDecodeError:
                                raise HTTPException(
                                    status_code=response.status, 
                                    detail=f"Gemini API error: {response_text[:200]}"
                                )
                        
            except asyncio.TimeoutError:
                if attempt < self.max_retries - 1:
                    delay = self.base_delay * (2 ** attempt)
                    logging.warning(f"Gemini API timeout, retrying in {delay}s (attempt {attempt + 1}/{self.max_retries})")
                    await asyncio.sleep(delay)
                    continue
                else:
                    raise HTTPException(
                        status_code=504, 
                        detail="Request timeout. The AI service is taking longer than expected. Please try again."
                    )
            
            except aiohttp.ClientError as e:
                if attempt < self.max_retries - 1:
                    delay = self.base_delay * (2 ** attempt)
                    logging.warning(f"Gemini API connection error: {str(e)}, retrying in {delay}s")
                    await asyncio.sleep(delay)
                    continue
                else:
                    raise HTTPException(
                        status_code=503, 
                        detail="Unable to connect to AI service. Please check your internet connection and try again."
                    )
            
            except Exception as e:
                # Don't retry for unexpected errors
                logging.error(f"Unexpected error in Gemini API call: {str(e)}")
                raise HTTPException(
                    status_code=500, 
                    detail=f"An unexpected error occurred during analysis: {str(e)}"
                )

    async def generate_workflow(self, document_content: str, workflow_type: str) -> List[WorkflowNode]:
        """Generate workflow from PRD/BRD using Gemini AI"""
        
        workflow_prompts = {
            "user_journey": """You are an expert product analyst. Carefully read the provided PRD/BRD document and extract the EXACT user journey described in the document.

Look for:
- Specific user stories mentioned in the document
- Exact user flows described
- Decision points where users make choices
- Specific features and how users interact with them
- Particular user roles mentioned
- Detailed step-by-step processes described

Create a user journey workflow that follows the EXACT steps described in the document, including proper decision points.

Generate a JSON array of workflow nodes with this structure:
[
  {
    "id": "unique_id",
    "type": "start|process|decision|end",
    "label": "Exact step from the document",
    "x": x_position,
    "y": y_position,
    "connections": [{"target": "next_node_id", "label": "Yes|No|Continue|etc"}]
  }
]

IMPORTANT RULES:
- Use "start" for entry points (rectangular)
- Use "process" for actions/steps (rectangular) 
- Use "decision" for choices/branches (diamond shape) with Yes/No or specific options
- Use "end" for completion points (oval)
- Use the EXACT terminology from the document
- Follow the SPECIFIC user flows mentioned 
- Include the ACTUAL features described
- Create proper decision branches with labels like "Yes", "No", "Success", "Failure"
- Position nodes: start at x:200, then increment by 300 (500, 800, 1100, etc)
- Use y:100 for main flow, y:300 for "No" branches, y:50 for "Yes" branches
- Each connection should have a "target" and "label"

Extract the real workflow from this specific document with proper decision branching.""",

            "service_blueprint": """You are an expert business process analyst. Carefully analyze the provided PRD/BRD document and extract the EXACT service delivery process described.

Look for:
- Backend processes mentioned in the document
- System interactions described
- Service delivery steps outlined
- Validation and approval steps
- Error handling processes
- Quality checks mentioned

Create a service blueprint that maps the SPECIFIC service processes described in the document with proper decision points.

Generate a JSON array of workflow nodes:
[
  {
    "id": "unique_id", 
    "type": "start|process|decision|end",
    "label": "Specific service step from document",
    "x": x_position,
    "y": y_position,
    "connections": [{"target": "next_node_id", "label": "Approved|Rejected|Continue|etc"}]
  }
]

Include proper decision diamonds for validation, approval, and error handling steps.""",

            "feature_flow": """You are an expert technical architect. Analyze the provided PRD/BRD document and extract the EXACT feature interactions and technical flow described.

Look for:
- Specific features mentioned and how they connect
- Technical integrations described
- Conditional logic and branching
- Feature dependencies outlined
- Error handling and validation steps

Create a feature flow that shows the SPECIFIC feature interactions described in the document with proper conditional branches.

Generate a JSON array of workflow nodes:
[
  {
    "id": "unique_id",
    "type": "start|process|decision|end", 
    "label": "Actual feature/component from document",
    "x": x_position,
    "y": y_position,
    "connections": [{"target": "next_node_id", "label": "Success|Error|Valid|Invalid|etc"}]
  }
]

Include proper decision points for feature validation and conditional logic."""
        }

        prompt = workflow_prompts.get(workflow_type, workflow_prompts["user_journey"])
        
        payload = {
            "contents": [
                {
                    "parts": [
                        {
                            "text": f"{prompt}\n\nDocument to analyze:\n\n{document_content}"
                        }
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.1,  # Lower temperature for more precise extraction
                "topK": 20,
                "topP": 0.8,
                "maxOutputTokens": 2048,
            }
        }
        
        headers = {
            "Content-Type": "application/json",
            "X-goog-api-key": self.api_key
        }
        
        # Use same retry logic as document analysis
        for attempt in range(self.max_retries):
            try:
                timeout_duration = 30 + (attempt * 15)
                
                async with aiohttp.ClientSession() as session:
                    async with session.post(
                        self.base_url, 
                        json=payload, 
                        headers=headers,
                        timeout=aiohttp.ClientTimeout(total=timeout_duration)
                    ) as response:
                        response_text = await response.text()
                        
                        if response.status == 200:
                            result = json.loads(response_text)
                            
                            if 'candidates' in result and len(result['candidates']) > 0:
                                content = result['candidates'][0]['content']['parts'][0]['text']
                                
                                try:
                                    # Clean up the content first - remove markdown code blocks
                                    cleaned_content = content.strip()
                                    if cleaned_content.startswith('```json'):
                                        cleaned_content = cleaned_content.replace('```json', '').replace('```', '').strip()
                                    elif cleaned_content.startswith('```'):
                                        cleaned_content = cleaned_content.replace('```', '').strip()
                                    
                                    # Try to parse as JSON array
                                    workflow_data = json.loads(cleaned_content)
                                    
                                    # Handle both array and object responses
                                    if isinstance(workflow_data, dict) and 'workflow' in workflow_data:
                                        workflow_data = workflow_data['workflow']
                                    elif isinstance(workflow_data, dict) and 'nodes' in workflow_data:
                                        workflow_data = workflow_data['nodes']
                                    
                                    # Convert to WorkflowNode objects with enhanced connection handling
                                    nodes = []
                                    for i, node_data in enumerate(workflow_data):
                                        if isinstance(node_data, dict):
                                            # Handle different connection formats
                                            connections = node_data.get('connections', [])
                                            connection_list = []
                                            
                                            for conn in connections:
                                                if isinstance(conn, dict):
                                                    # New format with labels: [{"target": "id", "label": "Yes"}]
                                                    target = conn.get('target', '')
                                                    if target:
                                                        connection_list.append(target)
                                                elif isinstance(conn, str):
                                                    # Simple format: ["id1", "id2"]
                                                    connection_list.append(conn)
                                            
                                            nodes.append(WorkflowNode(
                                                id=node_data.get('id', f'node_{i}'),
                                                type=node_data.get('type', 'process'),
                                                label=node_data.get('label', f'Step {i+1}'),
                                                x=node_data.get('x', 200 + (i * 300)),
                                                y=node_data.get('y', 100),
                                                connections=connection_list
                                            ))
                                    
                                    # If we got valid nodes, return them
                                    if nodes and len(nodes) > 2:
                                        return nodes
                                    else:
                                        # Generate better fallback based on document content
                                        return self._create_smart_fallback_workflow(document_content, workflow_type)
                                        
                                except json.JSONDecodeError:
                                    # Try to extract JSON from within the content
                                    import re
                                    json_match = re.search(r'\[.*\]', content, re.DOTALL)
                                    if json_match:
                                        try:
                                            workflow_data = json.loads(json_match.group())
                                            nodes = []
                                            for i, node_data in enumerate(workflow_data):
                                                if isinstance(node_data, dict):
                                                    # Handle different connection formats
                                                    connections = node_data.get('connections', [])
                                                    connection_list = []
                                                    
                                                    for conn in connections:
                                                        if isinstance(conn, dict):
                                                            # New format with labels: [{"target": "id", "label": "Yes"}]
                                                            target = conn.get('target', '')
                                                            if target:
                                                                connection_list.append(target)
                                                        elif isinstance(conn, str):
                                                            # Simple format: ["id1", "id2"]
                                                            connection_list.append(conn)
                                                    
                                                    nodes.append(WorkflowNode(
                                                        id=node_data.get('id', f'node_{i}'),
                                                        type=node_data.get('type', 'process'),
                                                        label=node_data.get('label', f'Step {i+1}'),
                                                        x=node_data.get('x', 100 + (i * 150)),
                                                        y=node_data.get('y', 100),
                                                        connections=connection_list
                                                    ))
                                            if nodes and len(nodes) > 2:
                                                return nodes
                                        except json.JSONDecodeError:
                                            pass
                                    
                                    # If all parsing fails, create smart fallback
                                    return self._create_smart_fallback_workflow(document_content, workflow_type)
                            else:
                                raise HTTPException(status_code=500, detail="No workflow generated")
                        
                        elif response.status == 503:
                            if attempt < self.max_retries - 1:
                                delay = self.base_delay * (2 ** attempt)
                                await asyncio.sleep(delay)
                                continue
                            else:
                                raise HTTPException(
                                    status_code=503,
                                    detail="AI service temporarily overloaded. Please try generating workflow again in a moment."
                                )
                        else:
                            raise HTTPException(
                                status_code=response.status,
                                detail=f"Workflow generation failed: {response_text[:200]}"
                            )
                            
            except Exception as e:
                if attempt < self.max_retries - 1:
                    delay = self.base_delay * (2 ** attempt)
                    await asyncio.sleep(delay)
                    continue
                else:
                    raise HTTPException(
                        status_code=500,
                        detail=f"Workflow generation failed: {str(e)}"
                    )

    def _create_smart_fallback_workflow(self, content: str, workflow_type: str) -> List[WorkflowNode]:
        """Create a smart fallback workflow based on document content with proper decision points"""
        
        # Analyze content for keywords and structure
        content_lower = content.lower()
        
        if workflow_type == "user_journey":
            # Create user journey based on document content
            if "e-commerce" in content_lower or "shopping" in content_lower or "cart" in content_lower:
                return [
                    WorkflowNode(id="start", type="start", label="User Opens App", x=200, y=100, connections=["browse"]),
                    WorkflowNode(id="browse", type="process", label="Browse Products", x=500, y=100, connections=["select"]),
                    WorkflowNode(id="select", type="process", label="Select Items", x=800, y=100, connections=["cart"]),
                    WorkflowNode(id="cart", type="process", label="Add to Cart", x=1100, y=100, connections=["auth_check"]),
                    WorkflowNode(id="auth_check", type="decision", label="User Logged In?", x=1400, y=100, connections=["payment", "login"]),
                    WorkflowNode(id="login", type="process", label="Login/Register", x=1400, y=300, connections=["payment"]),
                    WorkflowNode(id="payment", type="process", label="Process Payment", x=1700, y=100, connections=["success_check"]),
                    WorkflowNode(id="success_check", type="decision", label="Payment Success?", x=2000, y=100, connections=["confirmation", "retry"]),
                    WorkflowNode(id="retry", type="process", label="Retry Payment", x=2000, y=300, connections=["payment"]),
                    WorkflowNode(id="confirmation", type="process", label="Order Confirmation", x=2300, y=100, connections=["end"]),
                    WorkflowNode(id="end", type="end", label="Order Complete", x=2600, y=100, connections=[])
                ]
            elif "social" in content_lower or "post" in content_lower or "media" in content_lower:
                return [
                    WorkflowNode(id="start", type="start", label="User Login", x=200, y=100, connections=["dashboard"]),
                    WorkflowNode(id="dashboard", type="process", label="View Dashboard", x=500, y=100, connections=["create"]),
                    WorkflowNode(id="create", type="process", label="Create Post", x=800, y=100, connections=["content_check"]),
                    WorkflowNode(id="content_check", type="decision", label="Content Valid?", x=1100, y=100, connections=["schedule", "edit"]),
                    WorkflowNode(id="edit", type="process", label="Edit Content", x=1100, y=300, connections=["content_check"]),
                    WorkflowNode(id="schedule", type="decision", label="Schedule or Publish?", x=1400, y=100, connections=["publish", "schedule_time"]),
                    WorkflowNode(id="schedule_time", type="process", label="Set Schedule Time", x=1400, y=50, connections=["publish"]),
                    WorkflowNode(id="publish", type="process", label="Publish Content", x=1700, y=100, connections=["end"]),
                    WorkflowNode(id="end", type="end", label="Post Published", x=2000, y=100, connections=[])
                ]
            else:
                return [
                    WorkflowNode(id="start", type="start", label="User Entry", x=200, y=100, connections=["discover"]),
                    WorkflowNode(id="discover", type="process", label="Discover Features", x=500, y=100, connections=["valid_check"]),
                    WorkflowNode(id="valid_check", type="decision", label="Valid Input?", x=800, y=100, connections=["interact", "error"]),
                    WorkflowNode(id="error", type="process", label="Show Error Message", x=800, y=300, connections=["discover"]),
                    WorkflowNode(id="interact", type="process", label="User Interaction", x=1100, y=100, connections=["complete"]),
                    WorkflowNode(id="complete", type="process", label="Complete Task", x=1400, y=100, connections=["end"]),
                    WorkflowNode(id="end", type="end", label="Task Complete", x=1700, y=100, connections=[])
                ]
        
        elif workflow_type == "service_blueprint":
            if "support" in content_lower or "ticket" in content_lower or "customer" in content_lower:
                return [
                    WorkflowNode(id="start", type="start", label="Service Request", x=200, y=100, connections=["validate"]),
                    WorkflowNode(id="validate", type="process", label="Validate Request", x=500, y=100, connections=["valid_check"]),
                    WorkflowNode(id="valid_check", type="decision", label="Request Valid?", x=800, y=100, connections=["route", "reject"]),
                    WorkflowNode(id="reject", type="process", label="Reject Request", x=800, y=300, connections=["notify_rejection"]),
                    WorkflowNode(id="notify_rejection", type="end", label="Rejection Sent", x=1100, y=300, connections=[]),
                    WorkflowNode(id="route", type="process", label="Route to Team", x=1100, y=100, connections=["priority"]),
                    WorkflowNode(id="priority", type="decision", label="High Priority?", x=1400, y=100, connections=["escalate", "resolve"]),
                    WorkflowNode(id="escalate", type="process", label="Escalate to Manager", x=1400, y=50, connections=["resolve"]),
                    WorkflowNode(id="resolve", type="process", label="Resolve Issue", x=1700, y=100, connections=["quality_check"]),
                    WorkflowNode(id="quality_check", type="decision", label="Quality Approved?", x=2000, y=100, connections=["notify", "rework"]),
                    WorkflowNode(id="rework", type="process", label="Rework Solution", x=2000, y=300, connections=["resolve"]),
                    WorkflowNode(id="notify", type="process", label="Notify Customer", x=2300, y=100, connections=["end"]),
                    WorkflowNode(id="end", type="end", label="Service Complete", x=2600, y=100, connections=[])
                ]
            else:
                return [
                    WorkflowNode(id="start", type="start", label="Service Trigger", x=200, y=100, connections=["validate"]),
                    WorkflowNode(id="validate", type="process", label="Input Validation", x=500, y=100, connections=["valid_check"]),
                    WorkflowNode(id="valid_check", type="decision", label="Input Valid?", x=800, y=100, connections=["process", "error"]),
                    WorkflowNode(id="error", type="end", label="Error Response", x=800, y=300, connections=[]),
                    WorkflowNode(id="process", type="process", label="Core Processing", x=1100, y=100, connections=["quality"]),
                    WorkflowNode(id="quality", type="decision", label="Quality Check Pass?", x=1400, y=100, connections=["deliver", "retry"]),
                    WorkflowNode(id="retry", type="process", label="Retry Processing", x=1400, y=300, connections=["process"]),
                    WorkflowNode(id="deliver", type="process", label="Service Delivery", x=1700, y=100, connections=["end"]),
                    WorkflowNode(id="end", type="end", label="Service Complete", x=2000, y=100, connections=[])
                ]
        
        elif workflow_type == "feature_flow":
            if "api" in content_lower or "integration" in content_lower:
                return [
                    WorkflowNode(id="start", type="start", label="API Request", x=200, y=100, connections=["auth"]),
                    WorkflowNode(id="auth", type="process", label="Authentication", x=500, y=100, connections=["auth_check"]),
                    WorkflowNode(id="auth_check", type="decision", label="Auth Valid?", x=800, y=100, connections=["validate", "unauthorized"]),
                    WorkflowNode(id="unauthorized", type="end", label="Unauthorized", x=800, y=300, connections=[]),
                    WorkflowNode(id="validate", type="process", label="Input Validation", x=1100, y=100, connections=["valid_check"]),
                    WorkflowNode(id="valid_check", type="decision", label="Valid Input?", x=1400, y=100, connections=["process", "bad_request"]),
                    WorkflowNode(id="bad_request", type="end", label="Bad Request", x=1400, y=300, connections=[]),
                    WorkflowNode(id="process", type="process", label="Business Logic", x=1700, y=100, connections=["integrate"]),
                    WorkflowNode(id="integrate", type="process", label="External Integration", x=2000, y=100, connections=["success_check"]),
                    WorkflowNode(id="success_check", type="decision", label="Integration Success?", x=2300, y=100, connections=["response", "error_response"]),
                    WorkflowNode(id="error_response", type="end", label="Error Response", x=2300, y=300, connections=[]),
                    WorkflowNode(id="response", type="end", label="Success Response", x=2600, y=100, connections=[])
                ]
            else:
                return [
                    WorkflowNode(id="start", type="start", label="Feature Trigger", x=200, y=100, connections=["input"]),
                    WorkflowNode(id="input", type="process", label="Input Processing", x=500, y=100, connections=["validation"]),
                    WorkflowNode(id="validation", type="decision", label="Input Valid?", x=800, y=100, connections=["logic", "error"]),
                    WorkflowNode(id="error", type="end", label="Validation Error", x=800, y=300, connections=[]),
                    WorkflowNode(id="logic", type="process", label="Core Logic", x=1100, y=100, connections=["decide"]),
                    WorkflowNode(id="decide", type="decision", label="Business Rules Met?", x=1400, y=100, connections=["output", "alternative"]),
                    WorkflowNode(id="alternative", type="process", label="Alternative Path", x=1400, y=300, connections=["output"]),
                    WorkflowNode(id="output", type="process", label="Generate Output", x=1700, y=100, connections=["end"]),
                    WorkflowNode(id="end", type="end", label="Feature Complete", x=2000, y=100, connections=[])
                ]
        
        # Default fallback with decisions
        return [
            WorkflowNode(id="start", type="start", label="Start", x=200, y=100, connections=["process1"]),
            WorkflowNode(id="process1", type="process", label="Initial Process", x=500, y=100, connections=["decision1"]),
            WorkflowNode(id="decision1", type="decision", label="Continue?", x=800, y=100, connections=["process2", "end_early"]),
            WorkflowNode(id="end_early", type="end", label="Early Exit", x=800, y=300, connections=[]),
            WorkflowNode(id="process2", type="process", label="Final Process", x=1100, y=100, connections=["end"]),
            WorkflowNode(id="end", type="end", label="Complete", x=1400, y=100, connections=[])
        ]

    def _create_fallback_workflow(self, content: str, workflow_type: str) -> List[WorkflowNode]:
        """Create a simple fallback workflow if AI parsing fails"""
        return self._create_smart_fallback_workflow(content, workflow_type)

# Initialize services
gemini_service = GeminiAIService(GEMINI_API_KEY)
auth_service = AuthService()

# Authentication routes
@api_router.post("/auth/register", response_model=TokenResponse)
async def register(user_data: UserCreate):
    """Register a new user"""
    try:
        # Check if user already exists
        raise HTTPException(status_code=503, detail="Database is disabled. Registration unavailable.")
        
        # Hash password
        hashed_password = auth_service.hash_password(user_data.password)
        
        # Create user
        user = User(
            email=user_data.email,
            name=user_data.name,
            hashed_password=hashed_password
        )
        
        # Save to database
        # DB disabled: user not saved
        
        # Create access token
        access_token = auth_service.create_access_token({"user_id": user.id})
        
        user_response = UserResponse(
            id=user.id,
            email=user.email,
            name=user.name,
            created_at=user.created_at,
            is_active=user.is_active
        )
        
        return TokenResponse(
            access_token=access_token,
            token_type="bearer",
            user=user_response
        )
        
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"Registration failed: {str(e)}")

@api_router.post("/auth/login", response_model=TokenResponse)
async def login(user_credentials: UserLogin):
    """Login user"""
    try:
        # Find user
        raise HTTPException(status_code=503, detail="Database is disabled. Login unavailable.")
        
        user = User(**user_data)
        
        # Verify password
        if not auth_service.verify_password(user_credentials.password, user.hashed_password):
            raise HTTPException(status_code=401, detail="Invalid email or password")
        
        # Create access token
        access_token = auth_service.create_access_token({"user_id": user.id})
        
        user_response = UserResponse(
            id=user.id,
            email=user.email,
            name=user.name,
            created_at=user.created_at,
            is_active=user.is_active
        )
        
        return TokenResponse(
            access_token=access_token,
            token_type="bearer",
            user=user_response
        )
        
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"Login failed: {str(e)}")

@api_router.get("/auth/me", response_model=UserResponse)
async def get_current_user_info(current_user: User = Depends(get_current_user)):
    """Get current user information"""
    return UserResponse(
        id=current_user.id,
        email=current_user.email,
        name=current_user.name,
        created_at=current_user.created_at,
        is_active=current_user.is_active
    )

# Basic routes
@api_router.get("/")
async def root():
    return {"message": "PRD/BRD AI Analysis API - Ready!"}

@api_router.post("/status", response_model=StatusCheck)
async def create_status_check(input: StatusCheckCreate):
    status_dict = input.dict()
    status_obj = StatusCheck(**status_dict)
    raise HTTPException(status_code=503, detail="Database is disabled. Status check unavailable.")
    return status_obj

@api_router.get("/status", response_model=List[StatusCheck])
async def get_status_checks():
    raise HTTPException(status_code=503, detail="Database is disabled. Status list unavailable.")

# Document analysis routes
@api_router.post("/analyze-document", response_model=DocumentAnalysisResponse)
async def analyze_document(
    request: DocumentAnalysisRequest,
    current_user: Optional[User] = Depends(get_current_user_optional)
):
    """Analyze a document using AI"""
    try:
        # Analyze with Gemini
        analysis_result = await gemini_service.analyze_document(
            request.document_content, 
            request.analysis_type
        )
        
        # Create analysis record
        analysis_record = DocumentAnalysis(
            document_content=request.document_content,
            analysis_result=analysis_result,
            document_length=len(request.document_content),
            analysis_type=request.analysis_type,
            user_id=current_user.id if current_user else None
        )
        
        # Database saving is disabled; just return the analysis result
        return DocumentAnalysisResponse(
            id=analysis_record.id,
            analysis_result=analysis_result,
            document_length=analysis_record.document_length,
            analysis_type=analysis_record.analysis_type,
            timestamp=analysis_record.timestamp,
            user_id=analysis_record.user_id
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Analysis failed: {str(e)}")

@api_router.post("/analyze-document-file")
async def analyze_document_file(
    file: UploadFile = File(...),
    analysis_type: str = Form("gap_analysis"),
    current_user: Optional[User] = Depends(get_current_user_optional)
):
    """Analyze uploaded document file"""
    try:
        # Read file content
        file_content = await file.read()
        
        # Parse document based on file type
        document_content = ""
        filename_lower = file.filename.lower()
        
        if filename_lower.endswith('.pdf'):
            document_content = DocumentParser.parse_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            document_content = DocumentParser.parse_docx(file_content)
        elif filename_lower.endswith(('.txt', '.md')):
            document_content = DocumentParser.parse_text(file_content)
        else:
            raise HTTPException(
                status_code=400, 
                detail="Unsupported file type. Please upload PDF, DOCX, TXT, or MD files."
            )
        
        if not document_content.strip():
            raise HTTPException(status_code=400, detail="No readable content found in the file")
        
        # Analyze with Gemini
        analysis_result = await gemini_service.analyze_document(
            document_content, 
            analysis_type
        )
        
        # Create analysis record
        analysis_record = DocumentAnalysis(
            document_content=document_content,
            analysis_result=analysis_result,
            document_length=len(document_content),
            analysis_type=analysis_type,
            user_id=current_user.id if current_user else None
        )
        
        # Database saving is disabled; just return the analysis result
        return {
            "id": analysis_record.id,
            "filename": file.filename,
            "analysis_result": analysis_result,
            "document_length": len(document_content),
            "analysis_type": analysis_type,
            "timestamp": analysis_record.timestamp,
            "user_id": analysis_record.user_id
        }
        
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"File analysis failed: {str(e)}")

@api_router.get("/analyses")
async def get_analyses(current_user: Optional[User] = Depends(get_current_user_optional)):
    """Get document analyses (user's own if authenticated, or public if not)"""
    try:
        # Build query
        query = {}
        if current_user:
            query["user_id"] = current_user.id
        
        raise HTTPException(status_code=503, detail="Database is disabled. Document analyses unavailable.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch analyses: {str(e)}")

@api_router.get("/analyses/{analysis_id}")
async def get_analysis(
    analysis_id: str,
    current_user: Optional[User] = Depends(get_current_user_optional)
):
    """Get specific analysis by ID"""
    try:
        # Build query
        query = {"id": analysis_id}
        if current_user:
            query["user_id"] = current_user.id
        
        raise HTTPException(status_code=503, detail="Database is disabled. Document analysis unavailable.")
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"Failed to fetch analysis: {str(e)}")

@api_router.post("/generate-workflow", response_model=WorkflowResponse)
async def generate_workflow(
    request: WorkflowRequest,
    current_user: Optional[User] = Depends(get_current_user_optional)
):
    """Generate workflow from PRD/BRD document"""
    try:
        # Generate workflow with Gemini
        workflow_nodes = await gemini_service.generate_workflow(
            request.document_content,
            request.workflow_type
        )
        
        # Create workflow record
        workflow_record = Workflow(
            document_content=request.document_content,
            workflow_nodes=workflow_nodes,
            workflow_type=request.workflow_type,
            document_length=len(request.document_content),
            user_id=current_user.id if current_user else None
        )
        
        # Save to database
        workflow_dict = workflow_record.dict()
        # Convert WorkflowNode objects to dicts for MongoDB
        workflow_dict["workflow_nodes"] = [node.dict() for node in workflow_nodes]
        raise HTTPException(status_code=503, detail="Database is disabled. Workflow save unavailable.")
        
        return WorkflowResponse(
            id=workflow_record.id,
            workflow_nodes=workflow_nodes,
            workflow_type=workflow_record.workflow_type,
            document_length=workflow_record.document_length,
            timestamp=workflow_record.timestamp,
            user_id=workflow_record.user_id
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Workflow generation failed: {str(e)}")

@api_router.get("/workflows")
async def get_workflows(current_user: Optional[User] = Depends(get_current_user_optional)):
    """Get workflows (user's own if authenticated)"""
    try:
        raise HTTPException(status_code=503, detail="Database is disabled. Workflow list unavailable.")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to fetch workflows: {str(e)}")

@api_router.get("/workflows/{workflow_id}")
async def get_workflow(
    workflow_id: str,
    current_user: Optional[User] = Depends(get_current_user_optional)
):
    """Get specific workflow by ID"""
    try:
        raise HTTPException(status_code=503, detail="Database is disabled. Workflow lookup unavailable.")
    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(status_code=500, detail=f"Failed to fetch workflow: {str(e)}")

@api_router.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow(),
        "gemini_api_configured": bool(GEMINI_API_KEY),
        "features": {
            "authentication": True,
            "document_parsing": ["PDF", "DOCX", "TXT", "MD"],
            "analysis_types": ["gap_analysis", "requirements_extraction", "summary"],
            "workflow_generation": ["user_journey", "service_blueprint", "feature_flow"]
        }
    }

# Include the router in the main app
app.include_router(api_router)

app.add_middleware(
    CORSMiddleware,
    allow_credentials=True,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

@app.on_event("shutdown")
async def shutdown_db_client():
    # Database client is disabled; nothing to close
    pass